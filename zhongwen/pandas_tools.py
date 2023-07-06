'輔助PANDAS工具'

class 使用者要求覆寫(Exception):
    pass

class 批號存在錯誤(Exception):
    def __init__(self, 批號, 批號欄名, 表格):
        self.批號 = 批號
        self.批號欄名 = 批號欄名
        self.表格 = 表格

    def __str__(self):
        return f'批號【{self.批號}】已存在，請指定覆寫=True！'

def 批次讀取(批號, 批號欄名, 表格, 資料庫, parse_dates=None):
    import pandas as pd
    sql = f'select * from {表格} where {批號欄名}=="{批號}"'
    return pd.read_sql_query(sql, 資料庫, index_col='index', parse_dates=parse_dates)

def 批次刪除(批號, 批號欄名, 表格, 資料庫):
    c = 資料庫.cursor()
    sql = f'delete from {表格} where {批號欄名}=="{批號}"'
    c.execute(sql)
    資料庫.commit()
    import logging
    logging.debug(f'批次刪除{批號}、{表格}……')

class 該批資料為空(Exception):pass
    

def 批次寫入(資料, 批號, 批號欄名, 表格, 資料庫, 覆寫=False):
    import logging
    import pandas as pd
    try:
        df = 批次讀取(批號, 批號欄名, 表格, 資料庫) 
        if df.empty: raise 該批資料為空(f'表格：{表格}、批號：{批號}')
        if 覆寫: 
            批次刪除(批號, 批號欄名, 表格, 資料庫)
            raise 使用者要求覆寫()
        else: raise 批號存在錯誤(批號, 批號欄名, 表格)
    except (pd.errors.DatabaseError, 使用者要求覆寫, 該批資料為空) as e:
        logging.debug(f'批次寫入{批號}、{表格}……')
        資料.to_sql(表格, 資料庫, if_exists='append')
        logging.debug(f'寫入成功！')

def 可顯示(查詢資料函數):
    '裝飾查詢資料函數，指名參數設為【顯示=True】，即將查詢結果以 html 顯示。'
    from functools import wraps
    @wraps(查詢資料函數)
    def wrapper(*args, 顯示=False, **kargs):
        try:
            display_rows = kargs['顯示筆數']
            del kargs['顯示筆數']
        except KeyError:
            display_rows = 100
        df = 查詢資料函數(*args, **kargs)
        if 顯示: show_html(df, 顯示筆數=display_rows)
        return df
    return wrapper

def 強調關鍵字(df, 欄位=[], 關鍵字=[]):
    for c in 欄位:
        for k in 關鍵字:
            df[c] = df[c].str.replace(k, f'<b style="background:pink">{k}</b>')
    return df

def 資料型態標準化(df):
    '依資料框欄位名稱自動將資料框之資料型態標準化'
    from zhongwen.date import 取日期
    from zhongwen.number import 轉數值
    for c in df.columns:
        p = f'.*日期'
        import re
        if re.match(p, c):
            df[c] = df[c].map(取日期)
        p = f'.*金額'
        if re.match(p, c):
            df[c] = df[c].map(轉數值)
    return df            

def 標準格式(整數欄位=[], 實數欄位=[], 百分比欄位=[], 日期欄位=[], 
             最大值顯著欄位=[], 隱藏欄位=[], 
             百分比漸層欄位=[], 
             採用民國日期格式=False
            ):
    def formatter(style):
        style.applymap(lambda r:'text-align:right')
        if 整數欄位:
            style.format('{:,.0f}', subset=整數欄位)
        if 百分比欄位:
            style.format('{:,.2%}', subset=百分比欄位)
        if 實數欄位:
            style.format('{:,.2f}', subset=實數欄位)
        if 日期欄位:
            style.format('{:%Y%m%d}', subset=日期欄位)
            if 採用民國日期格式:
                from zhongwen.date import 民國日期
                style.format(民國日期, subset=日期欄位)
        if 最大值顯著欄位:
            # style.highlight_max(最大值顯著欄位)
            style.background_gradient(axis=0, cmap='RdYlGn'
                                     ,subset=最大值顯著欄位)
        if 百分比漸層欄位:
            style.background_gradient(axis=0, cmap='RdYlGn', vmax=1, vmin=-1
                                      ,subset=百分比漸層欄位
                                      )
        if 隱藏欄位:
            style.hide(隱藏欄位, axis=1) # hide index
        tr_hover = {
            'selector': 'tr:hover',
            'props':[('background-color', '#ffffb3'), ('border-style', 'dotted')]
        }
        style.set_table_styles([tr_hover], overwrite=False)
        return style
    return formatter

def show_html(df, 無格式=False
             ,整數欄位=[], 實數欄位=[]
             ,百分比欄位=[] ,日期欄位=[] ,隱藏欄位=[]
             ,百分比漸層欄位=[]
             ,最大值顯著欄位=[]
             ,顯示筆數=100, 採用民國日期格式=False, 標題=None
             ):
    import pandas as pd
    if isinstance(df, pd.Series):
        df = df.to_frame()
    df = df.head(顯示筆數)
    if not 無格式:
        from warnings import warn
        try:
            return 自動格式(df
                           ,整數欄位
                           ,實數欄位
                           ,百分比欄位
                           ,日期欄位
                           ,隱藏欄位
                           ,百分比漸層欄位
                           ,最大值顯著欄位
                           ,顯示筆數
                           ,採用民國日期格式
                           ,顯示=True
                           )
        except (ValueError, KeyError, TypeError) as e: 
            import traceback
            trace = traceback.format_exc()
            warn(f'發生例外：{trace}')
    from pathlib import Path
    html = Path.home() / 'TEMP' / 'output.html'
    df.to_html(html)
    import os
    os.system(f'start {html}')

def 自動格式(df, 整數欄位=[] ,實數欄位=[], 百分比欄位=[]
            ,日期欄位=[] ,隱藏欄位=[]
            ,百分比漸層欄位=[], 最大值顯著欄位=[]
            ,顯示筆數=100, 採用民國日期格式=False, 顯示=None
            ,除錯提示=False, 顯示提示=False
            ):
    if 顯示筆數:
        df = df[:顯示筆數]
    columns = df.columns
    from zhongwen.number import 轉數值
    import pandas as pd
    import numpy as np
    for c in df.columns:
        pat = '^.*述|借貸$'
        import re
        if re.match(pat, c):
            continue
        pat = '^.*(數|金額|損益|股利|累計|差異|期末|負債|營收|\(元\))|本益比|成本|支出|存入|現值|借券|餘額|借|貸$'
        if re.match(pat, c):
            if np.issubclass_(df[c].dtype.type, np.integer) and not c in 隱藏欄位:
                try:
                    整數欄位.append(c)
                except AttributeError:
                    整數欄位 = [c]
        pat = '^現金轉換天數|股價|配息$'
        if re.match(pat, c):
            if df[c].dtype == float and not c in 隱藏欄位: 
                try:
                    實數欄位.append(c)
                except AttributeError:
                    實數欄位 = [c]
        pat = '^.+(率|比例|比|\(%\))$'
        if re.match(pat, c) and c not in ['本益比']:
            if df[c].dtype == float and not c in 隱藏欄位: 
                try:
                    百分比欄位.append(c)
                except AttributeError:
                    百分比欄位 = [c]
        pat = '.*日期.*'
        if re.match(pat, c):
            if df[c].dtype == "datetime64[ns]" and not c in 隱藏欄位:
                try:
                    日期欄位.append(c)
                except AttributeError:
                    日期欄位 = [c]
    if 整數欄位: 整數欄位 = [*set(整數欄位)]
    if 實數欄位: 實數欄位 = [*set(實數欄位)]
    if 百分比欄位: 百分比欄位 = [*set(百分比欄位)]
    from itertools import chain
    n = list(chain.from_iterable([l for l in [整數欄位, 實數欄位] if isinstance(l, list)]))
    if n: 最大值顯著欄位.extend(n)
    最大值顯著欄位 = [*set(最大值顯著欄位)]
    百分比漸層欄位.extend(百分比欄位)
    import logging
    logging.debug(f'整數欄位：{整數欄位!r}')
    logging.debug(f'實數欄位：{實數欄位!r}')
    logging.debug(f'百分比欄位：{百分比欄位!r}')
    logging.debug(f'百分比漸層欄位：{百分比漸層欄位!r}')
    logging.debug(f'最大值顯著欄位：{最大值顯著欄位!r}')
    logging.debug(f'日期欄位：{日期欄位!r}')
    logging.debug(f'隱藏欄位：{隱藏欄位!r}')
    s = df.style.pipe(標準格式(整數欄位, 實數欄位, 百分比欄位
                              ,日期欄位
                              ,百分比漸層欄位
                              ,隱藏欄位
                              ,最大值顯著欄位
                              ,採用民國日期格式=採用民國日期格式
                              ))
    tp = df.copy()
    for c in df.columns: tp[c] = c
    
    if 除錯提示:
        from pathlib import Path
        html = Path.home() / 'TEMP' / 'output.html'
        tp.to_html(html)
        import os
        os.system(f'start {html}')

    if 顯示提示:
        s = s.set_tooltips(tp)

    if 顯示:
        from pathlib import Path
        html = Path.home() / 'TEMP' / 'showdf.html'
        s.to_html(html, encoding='utf8', doctype_html=True)
        # s.to_html(html)
        import os
        os.system(f'start {html}')
    return s

def read_csv(*args, **kwargs):
    '編碼錯誤預設使用 replace。'
    from pathlib import Path
    fn = args[0]
    if isinstance(fn, Path):
        fn = str(fn)
    with open(fn, encoding=kwargs['encoding'], errors='replace') as f:
        import pandas as pd
        df = pd.read_csv(f, *args[1:], **kwargs)
        return df

def read_fwf(*args, **kwargs): 
    import pandas as pd
    mbyte_columns = kwargs['mbyte_columns']
    names = kwargs['names']
    encoding = kwargs['encoding']
    del kwargs['mbyte_columns']
    del kwargs['encoding']
    df = pd.read_fwf(args[0], header=None
                    ,converters={h:str for h in names}
                    ,encoding = 'cp437' # 利用cp437剛好用1位元編碼特性，表達ByteString
                    ,**kwargs
                    )
    def convert(cp437:str):
        try:
            return cp437.encode('cp437').decode(encoding, errors='replace') 
        except AttributeError:
            return cp437
    for c in mbyte_columns: 
        df[c] = df[c].map(convert)
    return df

def read_docx_tables(filename, tab_id=None, **kwargs):
    from warnings import warn
    warn("警告：【read_docx_tables】函數將廢棄，請改用【read_docx】函數！")
    return read_docx_tables(filename, tab_id, **kwargs)

def read_docx(filename, tab_id=None, **kwargs):
    """
    讀取 Word 文件(.docx)表格至 Pnadas DataFrame
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """
    from docx import Document
    def read_docx_tab(tab, **kwargs):
        import io
        import csv
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        import pandas as pd
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise

def 分割鏈(鏈, 長度):
    from itertools import islice
    it = iter(鏈)
    size = 長度
    chunk = list(iter(lambda: list(islice(it, size)), []))
    return chunk
